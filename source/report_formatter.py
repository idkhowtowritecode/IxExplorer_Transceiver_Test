from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx2pdf import convert
from datetime import datetime, timedelta
from constants import *
from datetime import datetime
import os, sys

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def timeFormat(msg):
    timeNow = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return f'[{timeNow}] {msg}'

class Keysight_Report(object):
    @timer
    def __init__(self, data):
        # === 設定資料 ===
        self.company_name = "Keysight Inc."
        self.output_docx = "Keysight_Report.docx"
        self.output_pdf = "Keysight_Report.pdf"
        self.pass_fail_status = "PASS"  # or "FAIL"
        self.keysight_logo_path = resource_path("keysight_logo.png")
        self.customer_logo_path = self.find_logo_file('customer_logo')
        # === 建立 Word 文件 ===
        self.doc = Document()
        self.change_properties()
        self.style()
        self.result_data = data
        self.checkPassFail()

        # === 封面頁（不顯示頁碼） ===
        self.section = self.doc.sections[0]

    def find_logo_file(self, base_path):
        extensions = ['.png', '.jpg', '.jpeg']
        for ext in extensions:
            full_path = base_path + ext
            if os.path.isfile(full_path):
                return full_path
        return None

    def change_properties(self):
        core_props = self.doc.core_properties
        core_props.author = "Ixia Tester"
        core_props.title = "Transceiver Test report"
        core_props.subject = "Keysight Ixia Automation Test"
        core_props.comments = "Created by Leadertech"

    def style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    @timer
    def addCover(self):
        self.section.start_type = WD_SECTION.NEW_PAGE
        self.section.different_first_page_header_footer = True

        # === 封面 Logo ===
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False

        # 左側 Logo
        if self.customer_logo_path:
            cell_left = table.cell(0, 0)
            paragraph_left = cell_left.paragraphs[0]
            run_left = paragraph_left.add_run()
            run_left.add_picture(self.customer_logo_path, width=Inches(2.5))

        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.5)

        # 右側 Logo
        cell_right = table.cell(0, 1)
        paragraph_right = cell_right.paragraphs[0]
        paragraph_right.alignment = 2  
        # 右對齊
        run_right = paragraph_right.add_run()
        run_right.add_picture(self.keysight_logo_path, width=Inches(2.5))

        self.doc.add_paragraph("Keysight Test Report", "Title")

        # === PASS/FAIL 彩色欄位 ===
        p = self.doc.add_paragraph()
        run = p.add_run(self.pass_fail_status)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_color = RGBColor(255, 255, 255)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), '92D050' if self.pass_fail_status == "PASS" else 'FF0000')
        p._element.get_or_add_pPr().append(shading_elm)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = font_color

        # === 測試基本資訊 ===
        test_mode = self.result_data.get('config', {}).get('test_mode', 'Framed')
        subtitle_table = self.doc.add_table(rows=3, cols=2)
        subtitle_table.rows[0].cells[0].text = "Report Header:"
        subtitle_table.rows[0].cells[0].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(18)
        subtitle_table.rows[0].cells[1].text = "Keysight Inc."
        subtitle_table.rows[0].cells[1].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(18)

        subtitle_table.rows[1].cells[0].text = "Report Date:"
        subtitle_table.rows[1].cells[0].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(18)
        subtitle_table.rows[1].cells[1].text = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
        subtitle_table.rows[1].cells[1].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(18)

        subtitle_table.rows[2].cells[0].text = "Type:"
        subtitle_table.rows[2].cells[0].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(18)
        subtitle_table.rows[2].cells[1].text = test_mode
        subtitle_table.rows[2].cells[1].paragraphs[0].alignment =WD_ALIGN_PARAGRAPH.LEFT
        subtitle_table.rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(18)
    @timer
    def addPageNumber(self):
        # === 第二節起用頁碼與頁首 ===
        section2 = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.section.different_first_page_header_footer = False

        # 讓頁碼從此頁起算為 1
        sectPr = section2._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)

        # 計算寬度
        page_width = section2.page_width
        margin_left = section2.left_margin
        margin_right = section2.right_margin
        content_width = page_width - margin_left - margin_right

        header = section2.header
        table = header.add_table(rows=1, cols=2, width=content_width)
        table.autofit = False
        cells = table.rows[0].cells

        # 左側：公司名稱
        cells[0].text = self.company_name
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 右側：頁碼欄位（Page {PAGE} of {=NUMPAGES-1}）
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Page ")

        # 插入 PAGE 欄位
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        fldText = OxmlElement('w:t')
        fldText.text = "1"
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.extend([fldChar_begin, instrText, fldChar_sep, fldText, fldChar_end])
        run.add_text(" of ")

        # --- 開始外層 { = ... } ---
        fldChar_begin_outer = OxmlElement('w:fldChar')
        fldChar_begin_outer.set(qn('w:fldCharType'), 'begin')

        instrText_outer = OxmlElement('w:instrText')
        instrText_outer.set(qn('xml:space'), 'preserve')
        instrText_outer.text = "= "

        # --- 內層欄位 { NUMPAGES } ---
        fldChar_begin_inner = OxmlElement('w:fldChar')
        fldChar_begin_inner.set(qn('w:fldCharType'), 'begin')

        instrText_inner = OxmlElement('w:instrText')
        instrText_inner.set(qn('xml:space'), 'preserve')
        instrText_inner.text = "NUMPAGES"

        fldChar_sep_inner = OxmlElement('w:fldChar')
        fldChar_sep_inner.set(qn('w:fldCharType'), 'separate')

        fldText_inner = OxmlElement('w:t')
        fldText_inner.text = "1"

        fldChar_end_inner = OxmlElement('w:fldChar')
        fldChar_end_inner.set(qn('w:fldCharType'), 'end')

        # --- 外層公式其餘文字（-1） ---
        instrText_tail = OxmlElement('w:instrText')
        instrText_tail.set(qn('xml:space'), 'preserve')
        instrText_tail.text = " - 1"

        fldChar_sep_outer = OxmlElement('w:fldChar')
        fldChar_sep_outer.set(qn('w:fldCharType'), 'separate')

        fldText_outer = OxmlElement('w:t')
        fldText_outer.text = "1"

        fldChar_end_outer = OxmlElement('w:fldChar')
        fldChar_end_outer.set(qn('w:fldCharType'), 'end')

        # --- Append 到 run 中 ---
        run._r.extend([
            fldChar_begin_outer,
            instrText_outer,

            fldChar_begin_inner,
            instrText_inner,
            fldChar_sep_inner,
            fldText_inner,
            fldChar_end_inner,

            instrText_tail,
            fldChar_sep_outer,
            fldText_outer,
            fldChar_end_outer
        ])

        # 加入底部邊框線
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            borders.append(bottom)
            tcPr.append(borders)
    
    def format_seconds_to_HHMMSS(self, seconds: int) -> str:
        td = timedelta(seconds=seconds)
        total_minutes = td.total_seconds() // 60
        hours = int(total_minutes // 60)
        minutes = int(total_minutes % 60)
        seconds = int(td.total_seconds() % 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    
    def checkPassFail(self):
        if 'berSummary' in self.result_data.keys():
            for port in self.result_data['berSummary']:
                if self.result_data['berSummary'][port]['ber'] != 'PASS':
                    self.pass_fail_status = "FAIL"
        else:
            for key in self.result_data.keys():
                if 'l2Summary' in key:
                    if (int(self.result_data[key]['tx'])==0 and int(self.result_data[key]['rx'])==0) or int(self.result_data[key]['loss'])>0:
                        self.pass_fail_status = "FAIL"
                elif 'fecSummary' in key:
                    for port in self.result_data[key]:
                        if self.result_data[key][port]['pre_fec'] != 'PASS' or self.result_data[key][port]['post_fec'] != 'PASS':
                            self.pass_fail_status = "FAIL"
        return 

    def set_horizontal_borders_only(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')

        # 下邊框開啟
        for border_name in ['bottom']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 粗細
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)

        tcPr.append(borders)

    @timer
    def addSummary(self):
        # === Summary 頁 ===
        self.doc.add_heading("SUMMARY", level=1)
        summary_table = self.doc.add_table(rows=8, cols=2)
        for row in summary_table.rows[:-1]:
            for cell in row.cells:
                self.set_horizontal_borders_only(cell)

        config_data = self.result_data.get('config', {})
        if config_data is {}:
            print('No Summary Data')
            return
        ixos_version = config_data.get('ixos_version', 'None')
        start_time = config_data.get('start_time', 'None')
        if start_time != 'None':
            start_time_split = start_time.split('-')
            start_time = f'{start_time_split[0][0:4]}/{start_time_split[0][4:6]}/{start_time_split[0][6:8]}  {start_time_split[1][0:2]}:{start_time_split[1][2:4]}:{start_time_split[1][4:6]}'
        end_time = config_data.get('end_time', 'None')
        if end_time != 'None':
            end_time_split = end_time.split('-')
            end_time = f'{end_time_split[0][0:4]}/{end_time_split[0][4:6]}/{end_time_split[0][6:8]}  {end_time_split[1][0:2]}:{end_time_split[1][2:4]}:{end_time_split[1][4:6]}'
        test_duration = self.format_seconds_to_HHMMSS(int(config_data.get('test_duration', '0')))
        """
        if 'berSummary' in self.result_data.keys():
            for port in self.result_data['berSummary']:
                if self.result_data['berSummary'][port]['ber'] != 'PASS':
                    self.pass_fail_status = "FAIL"
        else:
            for key in self.result_data.keys():
                if 'l2Summary' in key:
                    if (int(self.result_data[key]['tx'])==0 and int(self.result_data[key]['rx'])==0) or int(self.result_data[key]['loss'])>0:
                        self.pass_fail_status = "FAIL"
                elif 'fecSummary' in key:
                    for port in self.result_data[key]:
                        if self.result_data[key][port]['pre_fec'] != 'PASS' or self.result_data[key][port]['post_fec'] != 'PASS':
                            self.pass_fail_status = "FAIL"
        """
        module_type = ""
        module_version = ""
        module_sn = ""
        for key in self.result_data.keys():
            if 'transceiverDOM' in key:
                for port in self.result_data[key]: 
                    _type = self.result_data[key][port]['transceiverTypeProperty'].replace('{','').replace('}','').strip()
                    _version = self.result_data[key][port]['revComplianceProperty'].replace('{','').replace('}','').strip()
                    _sn = self.result_data[key][port]['serialNumberProperty'].replace('{','').replace('}','').strip()
                    module_type += f'{port} - {_type}\n'
                    module_version += f'{port} - {_version}\n'
                    module_sn += f'{port} - {_sn}\n'
                module_type = module_type.strip()
                module_version = module_version.strip()
                module_sn = module_sn.strip()
                break
        summary_data = [
            ("IxOS Version", ixos_version),
            ("Test Start Time", start_time),
            ("Test End Time", end_time),
            ("Test Duration", test_duration),
            ("Pass/Fail Verdict", self.pass_fail_status),
            ("Module Type", module_type),
            ("Module Version", module_version),
            ("Serial Number", module_sn)
        ]
        for i, (k, v) in enumerate(summary_data):
            row = summary_table.rows[i].cells
            row[0].text = k
            row[1].text = v
    
    def changeTableFont(self, table, fontSize=Pt(8), fontType='Calibri', aligment=WD_ALIGN_PARAGRAPH.LEFT):
        for row in table.rows:
            for cell in row.cells:
                for pg in cell.paragraphs:
                    for run in pg.runs:
                        run = pg.runs[0]
                        run.font.size = fontSize
                        run.font.name = fontType
                    pg.alignment = aligment
        return
    
    def changeCellColor(self, cell, fill_color='D9D9D9', font_color='FFFFFF'):
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run()
        run.font.bold = True
        
        r = int(font_color[0:2], 16)
        g = int(font_color[2:4], 16)
        b = int(font_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)  # Light gray
        tc_pr.append(shd)
    @timer
    def addDOM(self):
        # === DOM 頁 ===
        self.doc.add_page_break()
        test_mode = self.result_data['config'].get('test_mode','None')
        if test_mode == 'None':
            ## self.doc.add_paragraph("No Transceiver DOM (Digital Optical Monitoring data available.")
            return
        elif test_mode == 'Unframed':
            self.doc.add_heading(f"Transceiver DOM (Digital Optical Monitoring) - BERT", level=1)
            self.addPerDOM('BERT')
        elif test_mode == "Framed":
            frameSizeList =  self.result_data['config'].get(f'test_frameSize_list', '64').split(' ')
            for frameSize in frameSizeList:
                self.doc.add_heading(f"Transceiver DOM (Digital Optical Monitoring) - {frameSize}", level=1)
                self.addPerDOM(frameSize)
                break
            return    
    @timer
    def addPerDOM(self, frameSize = '64'):
        dom_data = self.result_data.get(f'transceiverDOM{frameSize}', [])
        if not dom_data:
            ## self.doc.add_paragraph("No DOM data available.")
            return
        _domIndex = 1
        for port in dom_data.keys():
            if _domIndex > 1:
                self.doc.add_page_break()
                self.doc.add_paragraph('\n')
            _domIndex += 1
            table = self.doc.add_table(rows=26, cols=11, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            
            table.cell(0,0).merge(table.cell(0,10))
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = port

            # Style header row
            for col in range(11):
                cell = table.cell(0, col)
                self.changeCellColor(cell)

            # Merge Cell
            for i in range(16):
                for j in range(5):
                    table.cell(i+1,2*j).merge(table.cell(i+1,2*j+1))

            # Change Column Width
            for i in range(11):
                table.columns[i].allow_autofit = False
                table.columns[i].width = Inches(0.6)
                table.rows[17].cells[i].width = Inches(0.6)
            table.columns[0].width = Inches(0.7)
            table.rows[17].cells[0].width = Inches(0.7)
            table.columns[1].width = Inches(0.7)
            table.rows[17].cells[1].width = Inches(0.7)
            table.columns[2].width = Inches(0.7)
            table.rows[17].cells[2].width = Inches(0.7)
            table.columns[3].width = Inches(0.7)
            table.rows[17].cells[3].width = Inches(0.7)
            table.columns[5].width = Inches(0.4)
            table.rows[17].cells[5].width = Inches(0.4)
            table.columns[6].width = Inches(0.5)
            table.rows[17].cells[6].width = Inches(0.5)
            #table.columns[7].width = Inches(0.6)
            #table.rows[17].cells[7].width = Inches(0.6)
            table.columns[8].width = Inches(0.4)
            table.rows[17].cells[8].width = Inches(0.4)
            table.columns[9].width = Inches(0.4)
            table.rows[17].cells[9].width = Inches(0.4)
            table.columns[10].width = Inches(0.4)
            table.rows[17].cells[10].width = Inches(0.4)

            # Change Row Height
            for row in table.rows:
                row.height = Pt(10)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            # Transceiver Info
            infoCaptions = [
                (0, ['Manufacturer','Type','MSA','Media Tech','Cable Lenth','Reported Power Class']),
                (2, ['Model','SN','Date Code(YYMMDDLL)','Media Connector','Identifier Type','Reported Max Power']),
                (4, ['Mfg Revision','Firmware Revision','Hardware Revision'])
            ]
            infoCaption_key = {
                'Manufacturer': 'manufacturer', 'Type': 'transceiverTypeProperty', 'MSA': 'revComplianceProperty', 'Media Tech': 'mediaTechProperty', 'Cable Lenth': 'cableLengthProperty', 'Reported Power Class': 'powerClassProperty',
                'Model': 'model','SN': 'serialNumber','Date Code(YYMMDDLL)': 'dateCodeProperty','Media Connector': 'mediaConnectorProperty','Identifier Type': 'identifierTypeProperty','Reported Max Power': 'maxPowerProperty',
                'Mfg Revision': 'mfgRevProperty', 'Firmware Revision': 'firmwareRevProperty', 'Hardware Revision': 'hardwareRevProperty'
            }
            _rowIndex = 1
            for (k,v) in infoCaptions:
                for i in range(len(v)):
                    cell = table.cell(i+_rowIndex,2*k)
                    cell.text = v[i]
                    # Value
                    cell = table.cell(i+_rowIndex,2*k+2)
                    cell.text = dom_data[port][infoCaption_key[v[i]]].replace('{','').replace('}','')
            for i in range(0,6):
                for j in range(0,5,2):
                    cell = table.cell(i+_rowIndex,2*j)
                    self.changeCellColor(cell)

            # Module Temperature and Voltage
            moduleCaptions = [
                (0, ['Modlue','Temperature','Supply Voltage']),
                (1, ['High Alarm']),
                (2, ['High Warn']),
                (3, ['Low Warn']),
                (4, ['Low Alarm'])
            ]
            moduleCaption_key = {
                'High Alarm': ['temperatureHighAlarm', 'supplyVolHighAlarm'],
                'High Warn': ['temperatureHighWarn', 'supplyVolHighWarn'],
                'Low Warn': ['temperatureLowWarn', 'supplyVolLowWarn'],
                'Low Alarm': ['temperatureLowAlarm', 'supplyVolLowAlarm']
            }
            _rowIndex = 8
            for (k,v) in moduleCaptions:
                for i in range(len(v)):
                    cell = table.cell(i+_rowIndex,2*k)
                    cell.text = v[i]
                    if v[i] == 'High Alarm':
                        self.changeCellColor(cell, 'F7C8D2')
                    elif v[i] == 'High Warn':
                        self.changeCellColor(cell, 'FFE4BC')
                    elif v[i] == 'Low Warn':
                        self.changeCellColor(cell, 'FFF4C2')
                    elif v[i] == 'Low Alarm':
                        self.changeCellColor(cell, '#C6E3F1')
                    # Value
                    if k > 0:
                        cell = table.cell(i+_rowIndex+1,2*k)
                        cell.text = dom_data[port][moduleCaption_key[v[i]][0]].replace('{','').replace('}','')
                        cell = table.cell(i+_rowIndex+2,2*k)
                        cell.text = dom_data[port][moduleCaption_key[v[i]][1]].replace('{','').replace('}','')
                    else:
                        self.changeCellColor(cell)
            # Lane Limits
            laneLimitCaptions = [
                (0, ['Lane Limits','Tx Optical Power','Rx Optical Power', 'Tx Bias Current']),
                (1, ['High Alarm']),
                (2, ['High Warn']),
                (3, ['Low Warn']),
                (4, ['Low Alarm'])
            ]
            laneLimitCaption_key = {
                'High Alarm': ['txOpticalPowerHighAlarm', 'rxOpticalPowerHighAlarm', 'txBiasCurrentHighAlarm'],
                'High Warn': ['txOpticalPowerHighWarn', 'rxOpticalPowerHighWarn', 'txBiasCurrentHighWarn'],
                'Low Warn': ['txOpticalPowerLowWarn', 'rxOpticalPowerLowWarn', 'txBiasCurrentLowWarn'],
                'Low Alarm': ['txOpticalPowerLowAlarm', 'rxOpticalPowerLowAlarm', 'txBiasCurrentLowAlarm']
            }
            _rowIndex = 12
            for (k,v) in laneLimitCaptions:
                for i in range(len(v)):
                    cell = table.cell(i+_rowIndex,2*k)
                    cell.text = v[i]
                    if v[i] == 'High Alarm':
                        self.changeCellColor(cell, 'F7C8D2')
                    elif v[i] == 'High Warn':
                        self.changeCellColor(cell, 'FFE4BC')
                    elif v[i] == 'Low Warn':
                        self.changeCellColor(cell, 'FFF4C2')
                    elif v[i] == 'Low Alarm':
                        self.changeCellColor(cell, '#C6E3F1')
                    # Value
                    if k > 0:
                        cell = table.cell(i+_rowIndex+1,2*k)
                        cell.text = dom_data[port][laneLimitCaption_key[v[i]][0]].replace('{','').replace('}','')
                        cell = table.cell(i+_rowIndex+2,2*k)
                        cell.text = dom_data[port][laneLimitCaption_key[v[i]][1]].replace('{','').replace('}','')
                        cell = table.cell(i+_rowIndex+3,2*k)
                        cell.text = dom_data[port][laneLimitCaption_key[v[i]][1]].replace('{','').replace('}','')
                    else:
                        self.changeCellColor(cell)
            # Optical Lanes
            opticalLaneCaptions = ['Host Lane','Port','Data Path State','Tx LOS','Tx CDR LOL','Media Lane','Tx Optical Power','Tx Bias Cureent','Rx Optical Power','Rx LOS','Rx CDR LOL']
            opticalLaneCaptions_key = {
                'Port': 'portName','Data Path State': 'hostDataPathState','Tx LOS': 'hostTxLos','Tx CDR LOL': 'hostTxCdrLol','Media Lane': 'hostToMediaLane',
                'Tx Optical Power': 'mediaTxOpticalPower','Tx Bias Cureent': 'mediaTxBiasCurrent','Rx Optical Power': 'mediaRxOpticalPower','Rx LOS': 'mediaRxLos','Rx CDR LOL': 'mediaRxCdrLol'
            }
            _rowIndex = 17
            for i in range(len(opticalLaneCaptions)):
                cell = table.cell(_rowIndex,i)
                cell.text = opticalLaneCaptions[i]
                self.changeCellColor(cell)
                if opticalLaneCaptions[i] == 'Host Lane':
                    for j in range(8):
                        cell = table.cell(_rowIndex+1+j,i)
                        cell.text = f'{j+1}'
                else:
                    for j in range(len(dom_data[port][opticalLaneCaptions_key[opticalLaneCaptions[i]]])):
                        cell = table.cell(_rowIndex+1+j,i)
                        cell.text = dom_data[port][opticalLaneCaptions_key[opticalLaneCaptions[i]]][j]

            # Font Setting

            _rowIndex = 0
            for row in table.rows:
                for cell in row.cells:
                    for pg in cell.paragraphs:
                        for run in pg.runs:
                            run = pg.runs[0]
                            run.font.size = Pt(6)
                            run.font.name = 'Calibri'
                        if _rowIndex > 7:
                            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _rowIndex += 1
    @timer
    def addAppSel(self):
        # Title
        self.doc.add_page_break()
        self.doc.add_heading(f"CMIS Applicatio Select", level=1)
        appSelData = self.result_data.get(f'appSel', {})
        if not appSelData:
            self.doc.add_paragraph("No CMIS Applicatio Select data available.")
            return
        '''
        infoCaptions = ['Current AppSel','Req App Sel','Config Status']
        infoCaptions_key = {
            'Current AppSel': 'appSelCurrentValueProperty',
            'Req App Sel': '',
           'Config Status': ''
        }'''
        _appSelIndex = 1
        for port in appSelData.keys():
            if _appSelIndex > 1:
                self.doc.add_page_break()
                self.doc.add_paragraph('\n')
            _appSelIndex += 1
            # App Sel State
            table = self.doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = port
            cell = table.cell(1,0)
            cell.text = 'Current AppSel'
            cell = table.cell(1,1)
            cell.text = appSelData[port].get('appSelCurrentValueProperty', '-')

            # App Sel
            self.doc.add_paragraph('Available Applications')
            appSels = appSelData[port].get('transceiverAppSel',[])
            table = self.doc.add_table(rows=len(appSels)+2, cols=12, style='Table Grid')
            cell = table.cell(0,0)
            cell.text = 'App'
            mergeCell = table.cell(0,1).merge(table.cell(0,6))
            mergeCell.text = 'Host Side'
            mergeCell = table.cell(0,7).merge(table.cell(0,11))
            mergeCell.text = 'Line Side Media'
            columnCaptions = ['Interface','Lane Speed (G bit/s)','Modulation','Lane Groups','Lanes','ID (HEex)','Interface','Lane Speed (G bit/s)','Lane Groups','Lanes','ID (Hex)']
            columnCaptions_key = ['hostElectricalIfName','hostLaneSpeed','hostModulation','hostLaneGroup','hostLaneCount','hostElectricalIfID','mediaIfName','mediaLaneSpeed','mediaLaneGroup','mediaLaneCount','mediaIfID']
            for i in range(len(columnCaptions)):
                cell = table.cell(1,i+1)
                cell.text = columnCaptions[i]
            for i in range(2):
                for j in range(12):
                    self.changeCellColor(table.cell(i,j))
            for i in range(len(appSels)):
                cell = table.cell(i+2,0)
                cell.text = f'{i+1}'
                for j in range(len(columnCaptions_key)):
                    cell = table.cell(i+2,j+1)
                    if j == 0:
                        cell.text = f'{i}'
                    cell.text = appSels[i][columnCaptions_key[j]]

            # Font Setting
            self.changeTableFont(table, Pt(6), aligment=WD_ALIGN_PARAGRAPH.CENTER)

            # App Sel Preview
            self.doc.add_paragraph('Preview of Auto selected applications')
            appSelPreviews = appSelData[port].get('transceiverAppSelPreview',{})
            table = self.doc.add_table(rows=len(appSelPreviews.keys())+1, cols=10, style='Table Grid')       
            table.allow_autofit  = False     
            table.autofit = False
            for i in range(9):
                table.columns[i].allow_autofit = False
                table.columns[i].width = Inches(0.5)
                table.rows[0].cells[i].width = Inches(0.5)
            table.columns[9].width = Inches(1.2)
            table.rows[0].cells[9].width = Inches(1.2)
            columnCaptions = ['Host Port Mode','Mod','Lane Groups', 'Lanes','AppSel','Link','Host Electrical','Lane Groups','Lanes','Note']
            columnCaptions_key = {
                'Host Port Mode': 'portMode',
                'Mod': 'portModulation',
                'Lane Groups': 'portFanouts',
                'Lanes': 'portLaneCount',
                'AppSel': 'appSelId',
                'Link': 'link',
                'Host Electrical': 'moduleHostElectricalIfName',
                'Module Lane Groups': 'moduleHostLaneGroup',
                'Module Lanes': 'moduleHostLaneCount',
                'Note': 'note'
            }
            for i in range(len(columnCaptions)):
                cell = table.cell(0,i)
                cell.text = columnCaptions[i]
                self.changeCellColor(cell)
            _index = 1
            for appSelPreview in appSelPreviews.keys():
                for j in range(len(columnCaptions)):
                    cell = table.cell(_index,j)
                    cell.text = appSelPreviews[appSelPreview][columnCaptions_key[columnCaptions[j]]]
                _index += 1
            # Font Setting
            self.changeTableFont(table, Pt(6), aligment=WD_ALIGN_PARAGRAPH.CENTER)
        return
    @timer
    def addBer(self):       
        ### BER Summary ### 
        self.doc.add_page_break()
        self.doc.add_heading(f"BERT Result Summary", level=1)
        ber_data = self.result_data.get(f'berSummary', {})
        if not ber_data:
            ## self.doc.add_paragraph("No BERT Result Summary data available.")
            return
        
        table = self.doc.add_table(rows=3, cols=len(ber_data.keys())+1, style='Table Grid')
        berCaptions = ['Bit Error Ratio','BER Threshold',"Pre-BER Pass/Fail Verdict"]
        berCaptions_key = {
            'BER Threshold': 'threshold',
            "Pre-BER Pass/Fail Verdict": 'ber'
        }
        for i in range(len(berCaptions)):
            cell = table.cell(i,0)
            cell.text = berCaptions[i]
            if 'Pass/Fail' in berCaptions[i]:
                self.changeCellColor(cell,font_color='FF0000')
            else:
                self.changeCellColor(cell)
        _portIndex = 1
        for port in ber_data.keys():
            cell = table.cell(0,_portIndex)
            cell.text = port
            self.changeCellColor(cell)
            for i in range(len(berCaptions)):
                if i == 0:
                    continue
                cell = table.cell(i,_portIndex)
                _text = f'{ber_data[port][berCaptions_key[berCaptions[i]]]}'
                cell.text = _text
                if _text == 'PASS':
                    self.changeCellColor(cell, fill_color='92D050')
                elif _text == "FAIL":
                    self.changeCellColor(cell, fill_color='FF0000')
            _portIndex +=1
        return
    @timer
    def addBerStats(self):
        ### BER Per Lane Statistics ### 
        self.doc.add_page_break()
        self.doc.add_heading(f"BERT Statistics", level=1)
        ber_data = self.result_data.get(f'berLane', {})
        if not ber_data:
            ## self.doc.add_paragraph("No BERT Statistics data available.")
            return
        
        for port in ber_data.keys():
            self.doc.add_paragraph(f'Port - {port}')
            berCaptions = ['Pattern Lock','Pattern Transmitted','Pattern Received','Bits Send','Bits Received','Bit Error Received','Bit Error Ratio']
            berCaptions_key = {
                'Pattern Lock': 'bertPatternLock',
                'Pattern Transmitted': 'bertPatternTransmitted',
                'Pattern Received': 'bertPatternReceived',
                'Bits Send': 'bertBitsSent',
                'Bits Received': 'bertBitsReceived',
                'Bit Error Received': 'bertBitErrorsReceived',
                'Bit Error Ratio': 'bertBitErrorRatio'
                }
            table = self.doc.add_table(rows=len(ber_data[port])+1, cols=len(berCaptions)+1, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = table.cell(0,0)
            cell.text = 'Physical Lane'
            self.changeCellColor(cell)
            for i in range(len(berCaptions)):
                cell = table.cell(0,i+1)
                cell.text = berCaptions[i]
                self.changeCellColor(cell)
            _berIndex = -1
            for ber in ber_data[port]:
                cell = table.cell(_berIndex+2,0)
                if _berIndex == -1:
                    cell.text = 'Totals'
                else:
                    cell.text = f'{_berIndex}'
                self.changeCellColor(cell)
                for i in range(len(berCaptions)):
                    cell = table.cell(_berIndex+2,i+1)
                    cell.text = ber_data[port][_berIndex+1][berCaptions_key[berCaptions[i]]]
                _berIndex += 1

            self.changeTableFont(table, Pt(7))
        return
    @timer
    def addPerFEC(self, frameSize='64'):
        ### FEC Summary ### 
        self.doc.add_page_break()
        self.doc.add_heading(f"FEC Result Summary - {frameSize}", level=1)
        fec_data = self.result_data.get(f'fecSummary{frameSize}', {})
        if not fec_data:
            ## self.doc.add_paragraph("No FEC Result Summary data available.")
            return

        
        table = self.doc.add_table(rows=5, cols=len(fec_data.keys())+1, style='Table Grid')
        fecCaptions = ['Frame Loss Ratio','Pre-FEC Standard','Pre-FEC Pass/Fail Verdict','Post-FEC Standard','Post-FEC Pass/Fail Verdict']
        fecCaptions_key = {
            'Pre-FEC Standard': "pre_fec_threshold",
            'Pre-FEC Pass/Fail Verdict': "pre_fec",
            'Post-FEC Standard': "post_fec_pre_fec_threshold",
            'Post-FEC Pass/Fail Verdict': "post_fec"
        }
        for i in range(len(fecCaptions)):
            cell = table.cell(i,0)
            cell.text = fecCaptions[i]
            if 'Pass/Fail' in fecCaptions[i]:
                self.changeCellColor(cell,font_color='FF0000')
            else:
                self.changeCellColor(cell)
        _portIndex = 1
        for port in fec_data.keys():
            cell = table.cell(0,_portIndex)
            cell.text = port
            self.changeCellColor(cell)
            for i in range(len(fecCaptions)):
                if i == 0:
                    continue
                cell = table.cell(i,_portIndex)
                _text = f'{fec_data[port][fecCaptions_key[fecCaptions[i]]]}'
                cell.text = _text
                if _text == 'PASS':
                    self.changeCellColor(cell, fill_color='92D050')
                elif _text == "FAIL":
                    self.changeCellColor(cell, fill_color='FF0000')
            _portIndex +=1
        return
    @timer
    def addFec(self):
        test_mode = self.result_data['config'].get(f'test_mode','None')
        if test_mode == 'None' or test_mode == 'Unframed':
            self.doc.add_page_break()
            self.doc.add_heading(f"FEC Result Summary", level=1)
            ## self.doc.add_paragraph("No FEC Result Summary data available.")
            return
        frameSizeList = self.result_data['config'].get(f'test_frameSize_list', '64').split(' ')
        for frameSize in frameSizeList:
            self.addPerFEC(frameSize)   
        return
    @timer
    def addPcs(self):
        ### PCS Lane Statistics ###
        test_mode = self.result_data['config'].get(f'test_mode','None')
        if test_mode == 'None' or test_mode == 'Unframed':
            self.doc.add_page_break()
            self.doc.add_heading(f"PCS Lane Statistics", level=1)
            ## self.doc.add_paragraph("No PCS Lane Statistics data available.")
            return
        frameSizeList = self.result_data['config'].get(f'test_frameSize_list', '64').split(' ')
        for frameSize in frameSizeList:
            self.addPerPcs(frameSize)   
        return
    @timer
    def addPerPcs(self, frameSize='64'):
        self.doc.add_page_break()
        self.doc.add_heading(f"PCS Lane Statistics - {frameSize}", level=1)
        pcs_data = self.result_data.get(f'pcsLane{frameSize}', {})
        if not pcs_data:
            ## self.doc.add_paragraph("No PCS Lane Statistics data available.")
            return
        _pcsIndex = 1
        for port in pcs_data.keys():
            if _pcsIndex > 1:
                self.doc.add_page_break()
                self.doc.add_paragraph('\n')
            _pcsIndex += 1
            self.doc.add_paragraph(f'Port - {port}')
            pcsCaptions = ['PCS Lane Marker Lock','PCS Lane Marker Map','Relative Lane Skew (ns)','PCS Lane Marker Error Count','FEC Symbol Error Count','FEC Correct Bit Count','FEC Symbol Error','FEC Correct Bit Rate']
            pcsCaptions_key = {
                'PCS Lane Marker Lock': 'pcsLaneMarkerLock',
                'PCS Lane Marker Map': 'pcsLaneMarkerMap',
                'Relative Lane Skew (ns)': 'relativeLaneSkew',
                'PCS Lane Marker Error Count': 'pcsLaneMarkerErrorCount',
                'FEC Symbol Error Count': 'fecSymbolErrorCount',
                'FEC Correct Bit Count': 'fecCorrectedBitsCount',
                'FEC Symbol Error': 'fecSymbolErrorRate',
                'FEC Correct Bit Rate': 'fecCorrectedBitRate'
                }
            table = self.doc.add_table(rows=len(pcs_data[port])+1, cols=len(pcsCaptions)+1, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = table.cell(0,0)
            cell.text = 'Physical Lane'
            self.changeCellColor(cell)
            for i in range(len(pcsCaptions)):
                cell = table.cell(0,i+1)
                cell.text = pcsCaptions[i]
                self.changeCellColor(cell)
            _pcsIndex = -1
            for ber in pcs_data[port]:
                cell = table.cell(_pcsIndex+2,0)
                if _pcsIndex == -1:
                    cell.text = 'Totals'
                else:
                    cell.text = f'{_pcsIndex}'
                self.changeCellColor(cell)
                for i in range(len(pcsCaptions)):
                    cell = table.cell(_pcsIndex+2,i+1)
                    cell.text = pcs_data[port][_pcsIndex+1][pcsCaptions_key[pcsCaptions[i]]]
                _pcsIndex += 1
            
            self.changeTableFont(table, fontSize=Pt(7))
        return
    @timer
    def addL2Sumamry(self):
        ### L2 Traffic Test Summary ###
        test_mode = self.result_data['config'].get(f'test_mode','None')
        self.doc.add_page_break()
        self.doc.add_heading(f"L2 Traffic Test Summary", level=1)
        if test_mode == 'None' or test_mode == 'Unframed':
            ## self.doc.add_paragraph("No L2 Traffic Test Summary data available.")
            return
        frameSizeList = self.result_data['config'].get(f'test_frameSize_list', '64').split(' ')
        l2Captions = ['Tx Count','Rx Count','Loss Count','Loss %']
        l2Captions_key = {
            'Tx Count':'tx',
            'Rx Count':'rx',
            'Loss Count':'loss',
            'Loss %':'loss%'
        }
        table = self.doc.add_table(rows=len(frameSizeList)+1, cols=len(l2Captions)+1, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.cell(0,0)
        cell.text = 'Frame Size'
        self.changeCellColor(cell)
        for i in range(len(l2Captions)):
            cell = table.cell(0,i+1)
            cell.text = l2Captions[i]
            self.changeCellColor(cell)
        for i in range(len(frameSizeList)):
            cell = table.cell(i+1,0)
            cell.text = frameSizeList[i]
            self.changeCellColor(cell)
        for i in range(len(frameSizeList)):
            l2sum_data = self.result_data[f'l2Summary{frameSizeList[i]}']
            for j in range(len(l2Captions)):
                cell = table.cell(i+1,j+1)
                cell.text = l2sum_data[l2Captions_key[l2Captions[j]]]
        self.changeTableFont(table, Pt(7))
        return
    @timer
    def addPortStats(self, frameSize='64'):
        self.doc.add_page_break()
        self.doc.add_heading(f"Port Statistics - {frameSize}", level=1)
        stats_data = self.result_data.get(f'portStats{frameSize}', {})
        if not stats_data:
            ## self.doc.add_paragraph("No Port Statistics data available.")
            return

        # Collect all fields from all ports
        first_port = next(iter(stats_data.values()))
        field_names = list(first_port.keys())

        # Create table: rows = fields, columns = ports
        table = self.doc.add_table(rows=len(field_names)+1, cols=len(stats_data)+1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        table.cell(0, 0).text = "Port Statistics"
        port_ids = list(stats_data.keys())
        for col, port_id in enumerate(port_ids, start=1):
            table.cell(0, col).text = f"Port {port_id}"

        # Data rows
        for row, field in enumerate(field_names, start=1):
            table.cell(row, 0).text = field
            for col, port_id in enumerate(port_ids, start=1):
                value = stats_data[port_id].get(field, "")
                table.cell(row, col).text = str(value)

        # Style header row
        for col in range(len(stats_data)+1):
            cell = table.cell(0, col)
            self.changeCellColor(cell)
        self.changeTableFont(table)
    @timer
    def addPortStatistics(self):
        test_mode = self.result_data['config'].get(f'test_mode','None')
        if test_mode == 'None':
            ## self.doc.add_paragraph("No Port Statistics data available.")
            return
        elif test_mode == 'Unframed':
            self.addPortStats('BERT')
        elif test_mode == "Framed":
            frameSizeList = self.result_data['config'].get(f'test_frameSize_list', '64').split(' ')
            for frameSize in frameSizeList:
                self.addPortStats(frameSize)
            return
    @timer
    def save(self):
        # === 儲存 Word 檔 ===
        result_dir = self.result_data['config'].get('result_dir', '')
        self.doc.save(f'{result_dir}/{self.output_docx}')
        print(f"✅ Word Export at {self.output_docx}")
    @timer
    def saveAsPdf(self):
        # === 轉換為 PDF ===
        result_dir = self.result_data['config'].get('result_dir', '')
        try:
            convert(f'{result_dir}/{self.output_docx}', f'{result_dir}/{self.output_pdf}')
            print(f"✅ PDF Export at {result_dir}/{self.output_pdf}")
        except Exception as e:
            print("⚠️ PDF Export Fail: ", e)
    @timer
    def generate(self):
        self.addCover()
        self.addPageNumber()
        self.addSummary()
        # self.addDOM()
        # self.addAppSel()
        # self.addBer()
        # self.addBerStats()
        # self.addFec()
        # self.addPcs()
        self.addL2Sumamry()
        self.addPortStatistics()
        self.save()
        self.saveAsPdf()
