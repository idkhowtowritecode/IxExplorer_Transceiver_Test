from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx2pdf import convert
from datetime import datetime
import os

class Keysight_Report(object):
    def __init__(self):
        # === 設定資料 ===
        self.company_name = "Keysight Inc."
        self.logo_path = "logo.png"
        self.dom_image = "dom.png"
        self.port_stats_image = "port_stats.png"
        self.output_docx = "Keysight_Report.docx"
        self.output_pdf = "Keysight_Report.pdf"
        self.pass_fail_status = "PASS"  # 或 "FAIL"
        self.section = self.doc.sections[0]

    def addCover(self):
        # === 封面頁（不顯示頁碼） ===
        self.section.start_type = WD_SECTION.NEW_PAGE
        self.section.different_first_page_header_footer = True

        # === 封面 Logo ===
        self.doc.add_picture(self.logo_path, width=Inches(2.5))
        self.doc.add_paragraph("Keysight Test Report", "Title")

        # === PASS/FAIL 彩色欄位 ===
        p = self.doc.add_paragraph()
        run = p.add_run(self.pass_fail_status)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_color = RGBColor(255, 255, 255)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:fill'), '92D050' if self.pass_fail_status == "PASS" else 'FF0000')
        p._element.get_or_add_pPr().append(shading_elm)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = font_color

        # === 測試基本資訊 ===
        self.doc.add_paragraph("\nReport Header: EXFO Inc.\nReport Date: " + datetime.now().strftime("%Y/%m/%d %H:%M:%S") + "\nType: EtherBERT - B1: Framed / B2: Framed")

    def addPageNumber(self):
        # === 第二節起用頁碼與頁首 ===
        section2 = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.section.different_first_page_header_footer = False

        # 讓頁碼從此頁起算為 1
        sectPr = section2._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)

        # 計算寬度
        page_width = section2.page_width
        margin_left = section2.left_margin
        margin_right = section2.right_margin
        content_width = page_width - margin_left - margin_right

        header = section2.header
        table = header.add_table(rows=1, cols=2, width=content_width)
        table.autofit = False
        cells = table.rows[0].cells

        # 左側：公司名稱
        cells[0].text = self.company_name
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 右側：頁碼欄位（Page {PAGE} of {=NUMPAGES-1}）
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Page ")

        # 插入 PAGE 欄位
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        fldText = OxmlElement('w:t')
        fldText.text = "1"
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.extend([fldChar_begin, instrText, fldChar_sep, fldText, fldChar_end])
        run.add_text(" of ")

        # --- 開始外層 { = ... } ---
        fldChar_begin_outer = OxmlElement('w:fldChar')
        fldChar_begin_outer.set(qn('w:fldCharType'), 'begin')

        instrText_outer = OxmlElement('w:instrText')
        instrText_outer.set(qn('xml:space'), 'preserve')
        instrText_outer.text = "= "

        # --- 內層欄位 { NUMPAGES } ---
        fldChar_begin_inner = OxmlElement('w:fldChar')
        fldChar_begin_inner.set(qn('w:fldCharType'), 'begin')

        instrText_inner = OxmlElement('w:instrText')
        instrText_inner.set(qn('xml:space'), 'preserve')
        instrText_inner.text = "NUMPAGES"

        fldChar_sep_inner = OxmlElement('w:fldChar')
        fldChar_sep_inner.set(qn('w:fldCharType'), 'separate')

        fldText_inner = OxmlElement('w:t')
        fldText_inner.text = "1"

        fldChar_end_inner = OxmlElement('w:fldChar')
        fldChar_end_inner.set(qn('w:fldCharType'), 'end')

        # --- 外層公式其餘文字（-1） ---
        instrText_tail = OxmlElement('w:instrText')
        instrText_tail.set(qn('xml:space'), 'preserve')
        instrText_tail.text = " - 1"

        fldChar_sep_outer = OxmlElement('w:fldChar')
        fldChar_sep_outer.set(qn('w:fldCharType'), 'separate')

        fldText_outer = OxmlElement('w:t')
        fldText_outer.text = "1"

        fldChar_end_outer = OxmlElement('w:fldChar')
        fldChar_end_outer.set(qn('w:fldCharType'), 'end')

        # --- Append 到 run 中 ---
        run._r.extend([
            fldChar_begin_outer,
            instrText_outer,

            fldChar_begin_inner,
            instrText_inner,
            fldChar_sep_inner,
            fldText_inner,
            fldChar_end_inner,

            instrText_tail,
            fldChar_sep_outer,
            fldText_outer,
            fldChar_end_outer
        ])

        # 加入底部邊框線
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            borders.append(bottom)
            tcPr.append(borders)

    def addSummary(self):
        # === Summary 頁 ===
        self.doc.add_heading("SUMMARY", level=1)

        summary_table = self.doc.add_table(rows=5, cols=2)
        summary_data = [
            ("Test Status", "Completed"),
            ("Pass/Fail Verdict", self.pass_fail_status),
            ("Start Time", "2025/02/17 11:05:25 AM"),
            ("Duration", "00:10:00"),
            ("Test Recovery", "0")
        ]
        for i, (k, v) in enumerate(summary_data):
            row = summary_table.rows[i].cells
            row[0].text = k
            row[1].text = v
    
    def addDOM(self):
        # === DOM 頁 ===
        self.doc.add_page_break()
        self.doc.add_heading("Transceiver Digital Optical Monitoring", level=1)
        self.doc.add_paragraph("DOM Screenshot:")
        if os.path.exists(self.dom_image):
            self.doc.add_picture(self.dom_image, width=Inches(6))

    def addPortStatistics(self):
        # === Port Statistics 頁 ===
        self.doc.add_page_break()
        self.doc.add_heading("Port Statistics", level=1)
        self.doc.add_paragraph("Port Statistics Screenshot:")
        if os.path.exists(self.port_stats_image):
            self.doc.add_picture(self.port_stats_image, width=Inches(6))

    def save(self):      
        # === 儲存 Word 檔 ===
        self.doc.save(self.output_docx)
    
    def saveAsPdf(self):
        # === 轉換為 PDF ===
        try:
            convert(self.output_docx, self.output_pdf)
            print(f"✅ PDF convert success：{self.output_pdf}")
        except Exception as e:
            print("⚠️ PDF convert fail：", e)
 
    def generate(self):
         # === 建立 Word 文件 ===
        self.doc = Document()
        self.addCover()
        self.addPageNumber()
        self.addSummary()
        self.addDOM()
        self.addPortStatistics()
        self.save()
        self.saveAsPdf()